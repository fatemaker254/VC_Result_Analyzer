"""
Builds .docx reports:
  - build_docx_report():        a single department's report (unchanged shape
                                 from before, now optionally titled with the
                                 department name + semester)
  - build_comparison_docx():    the cross-department comparison report
                                 (Table 1 / Table 2 chart / Top 3 / Table 3 /
                                 SGPA chart), matching the client's reference
                                 tables.
"""

import os
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def _add_metric_row(doc, label, count, percent):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(f"{count} students ({percent}%)")


def _add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        pass  # style not available in this docx template; default grid still works
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return table


# ---------------------------------------------------------------------------
# Per-department report
# ---------------------------------------------------------------------------

def build_docx_report(result: dict, detection: dict, chart_dir: str, out_path: str,
                       department_name: str = None, semester: str = None):
    doc = Document()

    title = "Student Performance Report (SGPA Analysis)"
    if department_name:
        title = f"{department_name} — Student Performance Report"
    if semester:
        title += f" (Sem {semester})"
    doc.add_heading(title, level=0)

    meta = doc.add_paragraph()
    meta.add_run(f"Generated: {datetime.now().strftime('%d %b %Y, %H:%M')}").italic = True

    doc.add_paragraph(
        f"Total students analyzed: {result['total_students']}    |    "
        f"Total papers detected: {result['total_papers']}"
    )

    _add_heading(doc, "Detected Columns", level=2)
    doc.add_paragraph(f"Paper / subject columns: {', '.join(result['paper_columns'])}")
    doc.add_paragraph(f"SGPA column used: {result['sgpa_column']}")
    if result.get("remarks_column"):
        doc.add_paragraph(f"Overall result column used: {result['remarks_column']}")

    _add_heading(doc, "Summary of Results", level=1)
    m = result["metrics"]
    _add_metric_row(doc, "Students with all papers cleared", m["all_cleared"]["count"], m["all_cleared"]["percent"])
    _add_metric_row(doc, "Students with at least one fail", m["fail"]["count"], m["fail"]["percent"])
    _add_metric_row(doc, "Students scoring SGPA < 5", m["sgpa_below_5"]["count"], m["sgpa_below_5"]["percent"])
    _add_metric_row(doc, "Students scoring SGPA 5 to 7.5", m["sgpa_5_to_7_5"]["count"], m["sgpa_5_to_7_5"]["percent"])
    _add_metric_row(doc, "Students scoring SGPA > 7.5", m["sgpa_above_7_5"]["count"], m["sgpa_above_7_5"]["percent"])

    _add_heading(doc, "Charts", level=1)

    chart_specs = [
        ("summary.png", "Overall Performance Summary"),
        ("pass_fail.png", "Pass vs Fail"),
        ("sgpa_distribution.png", "SGPA Distribution"),
    ]

    for filename, caption in chart_specs:
        img_path = os.path.join(chart_dir, filename)
        if not os.path.exists(img_path):
            continue
        _add_heading(doc, caption, level=2)
        doc.add_picture(img_path, width=Inches(5.5))
        last_p = doc.paragraphs[-1]
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(9)

    doc.save(out_path)
    return out_path


# ---------------------------------------------------------------------------
# Cross-department comparison report
# ---------------------------------------------------------------------------

def build_comparison_docx(batch_meta: dict, chart_dir: str, out_path: str):
    doc = Document()

    semester = batch_meta.get("semester") or ""
    title = "Department-wise Comparison Report"
    if semester:
        title += f" — Sem {semester}"
    doc.add_heading(title, level=0)

    meta = doc.add_paragraph()
    meta.add_run(f"Generated: {datetime.now().strftime('%d %b %Y, %H:%M')}").italic = True

    comparison = batch_meta["comparison"]
    rows = comparison["rows"]
    agg = comparison["aggregate"]

    _add_heading(doc, "Table 1: Aggregate and Subject-wise Result", level=1)
    headers1 = ["Dept/Subject", "Students Appeared", "Cleared", "Not Cleared", "Cleared %", "Not Cleared %"]
    table1_rows = [
        [r["department_name"], r["students_appeared"], r["cleared"], r["not_cleared"],
         f'{r["cleared_pct"]:.2f}', f'{r["not_cleared_pct"]:.2f}']
        for r in rows
    ]
    table1_rows.append([
        "AGGREGATE", agg["students_appeared"], agg["cleared"], agg["not_cleared"],
        f'{agg["cleared_pct"]:.2f}', f'{agg["not_cleared_pct"]:.2f}'
    ])
    _add_table(doc, headers1, table1_rows)

    doc.add_paragraph()
    _add_heading(doc, "Table 2: Subject-wise Students' Performance", level=1)
    img1 = os.path.join(chart_dir, "dept_comparison.png")
    if os.path.exists(img1):
        doc.add_picture(img1, width=Inches(6.0))
        last_p = doc.paragraphs[-1]
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if comparison.get("top3"):
        _add_heading(doc, "Top 3 Departments (by % cleared)", level=2)
        for i, t in enumerate(comparison["top3"], start=1):
            doc.add_paragraph(
                f'{t["department_name"]} — {t["cleared"]} out of {t["students_appeared"]} '
                f'students @ {t["cleared_pct"]:.2f}%',
                style="List Number",
            )

    doc.add_paragraph()
    _add_heading(doc, "Table 3: Students' Performance in terms of SGPA", level=1)
    headers3 = ["Dept/Subject", "Total Appeared", "SGPA<5 No.", "SGPA<5 %",
                "SGPA 5-7.5 No.", "SGPA 5-7.5 %", "SGPA>7.5 No.", "SGPA>7.5 %"]
    table3_rows = [
        [r["department_name"], r["students_appeared"],
         r["sgpa_below_5"]["count"], f'{r["sgpa_below_5"]["percent"]:.2f}',
         r["sgpa_5_to_7_5"]["count"], f'{r["sgpa_5_to_7_5"]["percent"]:.2f}',
         r["sgpa_above_7_5"]["count"], f'{r["sgpa_above_7_5"]["percent"]:.2f}']
        for r in rows
    ]
    table3_rows.append([
        "AGGREGATE", agg["students_appeared"],
        agg["sgpa_below_5"]["count"], f'{agg["sgpa_below_5"]["percent"]:.2f}',
        agg["sgpa_5_to_7_5"]["count"], f'{agg["sgpa_5_to_7_5"]["percent"]:.2f}',
        agg["sgpa_above_7_5"]["count"], f'{agg["sgpa_above_7_5"]["percent"]:.2f}',
    ])
    _add_table(doc, headers3, table3_rows)

    doc.add_paragraph()
    img2 = os.path.join(chart_dir, "sgpa_comparison.png")
    if os.path.exists(img2):
        _add_heading(doc, "SGPA-wise Department Performance", level=2)
        doc.add_picture(img2, width=Inches(6.0))
        last_p = doc.paragraphs[-1]
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if batch_meta.get("file_errors"):
        doc.add_paragraph()
        _add_heading(doc, "Files Skipped", level=2)
        for fe in batch_meta["file_errors"]:
            doc.add_paragraph(f'{fe.get("file")}: {fe.get("error")}', style="List Bullet")

    doc.save(out_path)
    return out_path
